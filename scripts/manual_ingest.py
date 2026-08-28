#!/usr/bin/env python3
"""Curated manual-candidate ingestion for the R&I Geopolitics Radar.

Manual files are *candidate/recovery sources*, never evidence by themselves.  The
pipeline parses bibliography-like records, compares them with the cumulative radar,
retrieves the cited source when possible, applies the same substantive A/B gate used
by the scanner, and records provenance/recall diagnostics without changing scan
timestamps.
"""
from __future__ import annotations

import argparse
import csv
import datetime as dt
import hashlib
import io
import json
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader

try:
    import yaml
except ModuleNotFoundError:  # pragma: no cover - optional until YAML is actually used
    yaml = None

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover - surfaced as a friendly runtime error
    Document = None

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts import scan_radar as sr

PROFILE_VERSION = "v17.12.2-bounded-manual-review"
SUPPORTED = {".docx", ".pdf", ".csv", ".json", ".yaml", ".yml", ".txt", ".md"}
URL_RE = re.compile(r"https?://[^\s<>()\[\]{}]+", re.I)
ITEM_RE = re.compile(r"^([A-Za-z]{1,3}\d{1,3})\s+(.*)$")
MANUAL_BLOCK_HEADING_RE = re.compile(r"^EU R&I in Geopolitical Context\s+[—–-]\s+Items found", re.I)
TYPE_RE = re.compile(r"\bType\s+(\d+)\s*[—–-]\s*([^.;]+)", re.I)
CURRENT_SECTION_RE = re.compile(r"^[A-Z]\.\s+")
NUMBERED_SECTION_RE = re.compile(r"^\d+\.\s+")
NUMBERED_SUBSECTION_RE = re.compile(r"^\d+\.\d+\s+")
SECONDARY_HINTS = (
    "as reported by", "secondary source", "cited via secondary", "exact title to be confirmed",
    "reported).", "reported by science|business", "cited secondhand", "reporting source",
    "collates the primary references",
)
VERIFICATION_HINTS = (
    "publication date within the window not yet confirmed",
    "exact publication date not established",
    "exact 2026 publication date not established",
    "confirm exact publication date",
    "verify issue assignment",
    "confirm exact publication date and landing url",
    "substitute the eib press release",
    "substitute the loket kennisveiligheid primary page",
    "currently sourced from news reporting only",
    "link is to the starting grants page",
    "confirm against the published erc work programme",
    "exact title to be confirmed",
    "cite the cern council resolution directly",
    "locate the iris paper",
    "cite reuters and the commission release directly",
    "window: check date",
)
OUTSIDE_WINDOW_HINTS = (
    "outside the strict window",
    "outside the 4-month window",
    "outside the four-month window",
    "outside window",
    "context only",
)
FRONTIER_CURRENT_SECTIONS = (
    "knowledge & people",
    "infrastructure & inputs",
    "conversion",
    "rules & institutions",
)

MATRIX_ROWS = {"K": "knowledge", "I": "infrastructure", "C": "conversion", "R": "rules"}
MATRIX_DIMENSIONS = set(MATRIX_ROWS.values())
MATRIX_QUADRANTS = {"A", "B", "C", "D"}


def clean(value: Any) -> str:
    return sr.clean_text(value)


def _norm(value: Any) -> str:
    return sr.normalized(clean(value))


def _file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def _canonical_url(url: str) -> str:
    url = clean(url)
    if not url:
        return ""
    try:
        p = urlparse(url)
        host = p.netloc.lower().removeprefix("www.")
        path = re.sub(r"/+", "/", p.path or "/").rstrip("/") or "/"
        # Query strings often carry stable EU document identifiers; retain them but drop fragments.
        query = p.query.lower()
        return f"{host}{path.lower()}" + (f"?{query}" if query else "")
    except Exception:
        return _norm(url)



def _is_generic_homepage(url: str) -> bool:
    """Return True for bare site roots that are too unspecific for exact-URL recovery."""
    raw = clean(url)
    if not raw:
        return True
    parsed = urlparse(raw if "://" in raw else "https://" + raw)
    return not (parsed.path or "").strip("/") and not parsed.query

def _doi(value: str) -> str:
    m = re.search(r"10\.\d{4,9}/[^\s?#<>()]+", clean(value), re.I)
    return m.group(0).rstrip(".,;)").lower() if m else ""


def _extract_date(text: str) -> str:
    """Return a deterministic ISO date without inventing the current day.

    Month-only and year-only bibliographies use the first day solely as a sortable
    representation; ``date_precision`` records the original precision separately.
    """
    months = {
        'jan': 1, 'january': 1, 'feb': 2, 'february': 2, 'mar': 3, 'march': 3,
        'apr': 4, 'april': 4, 'may': 5, 'jun': 6, 'june': 6, 'jul': 7, 'july': 7,
        'aug': 8, 'august': 8, 'sep': 9, 'september': 9, 'oct': 10, 'october': 10,
        'nov': 11, 'november': 11, 'dec': 12, 'december': 12,
    }
    m = re.search(r"\b(\d{1,2})\s+(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+(20\d{2})\b", text, re.I)
    if m:
        return dt.date(int(m.group(3)), months[m.group(2).lower()], int(m.group(1))).isoformat()
    m = re.search(r"\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)(?:/\w+)?\s+(20\d{2})\b", text, re.I)
    if m:
        return dt.date(int(m.group(2)), months[m.group(1).lower()], 1).isoformat()
    # Common form: "2026 (online April 2026)" or "2026 (April; referenced ...)".
    m = re.search(r"\b(20\d{2})\b[^)]{0,40}\b(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\b", text, re.I)
    if m:
        return dt.date(int(m.group(1)), months[m.group(2).lower()], 1).isoformat()
    m = re.search(r"\b(20\d{2})\b", text)
    return f"{m.group(1)}-01-01" if m else ""


def _date_precision(text: str) -> str:
    if re.search(r"\b\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+20\d{2}\b", text, re.I):
        return 'day'
    if re.search(r"(?:\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*(?:/\w+)?\s+20\d{2}\b|\b20\d{2}\b[^)]{0,40}\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\b)", text, re.I):
        return 'month'
    if re.search(r"\b20\d{2}\b", text):
        return 'year'
    return ''


def _extract_title(citation: str, mode: str = "current") -> str:
    c = clean(citation)
    if not c:
        return ""
    if mode == "forthcoming":
        return clean(re.split(r"\s+[–—-]\s+", c, maxsplit=1)[0])
    if mode == "context":
        # Context lines are often "Institution, Title, date" rather than full bibliographies.
        c2 = re.sub(r"^(?:European Commission|European Parliament|Council of the European Union|European Court of Auditors|Science\|Business),\s*", "", c, flags=re.I)
        return clean(re.split(r",\s*(?:\d{1,2}\s+)?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d{2})", c2, maxsplit=1, flags=re.I)[0])
    # Most curated bibliography records have "(date). Title. Source".
    # Some recovery notes put "as cited in ..." between the date and the title.
    # Strip that bibliographic provenance clause before title extraction.
    cited = re.search(r"\)\s*,\s*(?:as cited in|as reported in)[^.]+\.\s*(.+)$", c, re.I)
    if cited:
        c = cited.group(1)
    m = re.search(r"\)\.\s*(.+)$", c)
    tail = m.group(1) if m else c
    # Avoid cutting titles at common geopolitical abbreviations such as "U.S.".
    parts = re.split(r"(?<!U\.S)(?<!U\.K)(?<!E\.U)\.\s+(?=[A-Z0-9])", tail, maxsplit=1)
    title = clean(parts[0])
    if len(title) < 12 and len(parts) > 1:
        title = clean(parts[1])
    return title.strip(" .")


def _extract_authors(citation: str) -> str:
    c = clean(citation)
    m = re.search(r"\s\((?:[^)]*20\d{2}[^)]*)\)\.\s*", c)
    return clean(c[:m.start()]) if m else ""


def _extract_source(citation: str, title: str) -> str:
    c = clean(citation)
    if title and title in c:
        tail = clean(c.split(title, 1)[1]).lstrip(". ")
        if tail:
            return clean(re.split(r"\.\s+", tail, maxsplit=1)[0])[:180]
    return ""


def _record_from_parts(*, rid: str, citation: str, url: str = "", note: str = "", section: str = "", mode: str = "current", source_file: str = "") -> dict[str, Any]:
    citation = clean(citation)
    note = clean(note)
    found_urls = URL_RE.findall(" ".join([citation, url, note]))
    url = clean(url or (found_urls[0] if found_urls else "")).rstrip(".,;)")
    title = _extract_title(citation, mode)
    status = "candidate"
    if mode == "forthcoming":
        status = "forthcoming_unpublished"
    elif mode == "context":
        status = "context_outside_primary_window"
    full_note = _norm(f"{citation} {note}")
    if status == "candidate" and any(x in full_note for x in OUTSIDE_WINDOW_HINTS):
        status = "context_outside_primary_window"
    secondary = any(x in full_note for x in SECONDARY_HINTS)
    verification_required = any(x in full_note for x in VERIFICATION_HINTS)
    cells = []
    for cell in re.findall(r"\b([KICR]-[ABCD])\b", f"{citation} {note}", re.I):
        cell = cell.upper()
        if cell not in cells:
            cells.append(cell)
    primary_cell = ""
    cell_text = f"{citation} {note}"
    # Curators may list several cells and mark a later one as primary, e.g.
    # "Cells: I-A, C-B (primary); I-C".  Prefer that explicit marker and only
    # fall back to the first listed cell when no marker is supplied.
    m = re.search(r"\b([KICR]-[ABCD])\s*\(primary\)", cell_text, re.I)
    if not m:
        m = re.search(r"Cells?:\s*([KICR]-[ABCD])", cell_text, re.I)
    if m:
        primary_cell = m.group(1).upper()
    type_match = TYPE_RE.search(f"{citation} {note}")
    curator_source_type = clean(type_match.group(2)) if type_match else ""
    verification_caveats = []
    if "window: check date" in full_note:
        verification_caveats.append("window_date_check")
    if secondary:
        verification_caveats.append("secondary_or_incomplete_source_hint")
    if verification_required and not verification_caveats:
        verification_caveats.append("curator_verification_required")
    return {
        "manual_id": rid,
        "title": title,
        "authors": _extract_authors(citation),
        "source": _extract_source(citation, title),
        "date": _extract_date(citation),
        "date_precision": _date_precision(citation),
        "url": url,
        "doi": _doi(f"{citation} {url}"),
        "citation": citation,
        "curator_note": note,
        "section": section,
        "manual_record_status": status,
        "manual_candidate_kind": "weak_signal" if mode == "weak_signal" else "substantive",
        "manual_secondary_hint": secondary,
        "manual_verification_required": verification_required,
        "verification_caveats": verification_caveats,
        "curator_source_type": curator_source_type,
        "curator_cells": cells,
        "curator_primary_cell": primary_cell,
        "curator_cell_mapping_status": "manual_hint_not_source_evidence" if cells else "",
        "source_file": source_file,
    }


def _parse_paragraph_records(paragraphs: Iterable[str], source_file: str) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    section = ""
    mode = "current"
    seq = {"forthcoming": 0, "context": 0, "generic": 0}
    manual_block_seen = False

    def flush() -> None:
        nonlocal current
        if current:
            records.append(_record_from_parts(source_file=source_file, **current))
            current = None

    for raw in paragraphs:
        p = clean(raw)
        if not p:
            continue
        if MANUAL_BLOCK_HEADING_RE.match(p):
            if manual_block_seen:
                # A curated DOCX may contain later rounds appended after the declared
                # batch. One ingestion call represents one batch, so stop at the next
                # top-level "Items found" heading instead of silently mixing rounds.
                flush()
                break
            manual_block_seen = True
            continue
        if CURRENT_SECTION_RE.match(p):
            flush(); section = p; mode = "current"; continue
        if NUMBERED_SUBSECTION_RE.match(p):
            flush(); section = p
            low = _norm(p)
            if p.startswith("5.") or "weak signal" in low:
                mode = "weak_signal"
            elif any(x in low for x in FRONTIER_CURRENT_SECTIONS):
                mode = "current"
            continue
        if NUMBERED_SECTION_RE.match(p):
            flush(); section = p
            low = _norm(p)
            if "expected" in low or "not yet published" in low:
                mode = "forthcoming"
            elif "essential" in low and "context" in low:
                mode = "context"
            elif "method" in low or "judgment calls" in low:
                break
            elif any(x in low for x in FRONTIER_CURRENT_SECTIONS):
                mode = "current"
            elif "weak signals" in low or re.match(r"^5\.\d+\s+", p):
                mode = "weak_signal"
            else:
                mode = "generic"
            continue
        m = ITEM_RE.match(p)
        if m and mode in {"current", "weak_signal"}:
            flush()
            current = {"rid": m.group(1), "citation": m.group(2), "url": "", "note": "", "section": section, "mode": mode}
            continue
        if mode in {"forthcoming", "context"} and not p.startswith("http"):
            flush()
            seq[mode] += 1
            prefix = "F" if mode == "forthcoming" else "X"
            records.append(_record_from_parts(rid=f"{prefix}{seq[mode]}", citation=p, section=section, mode=mode, source_file=source_file))
            continue
        if current:
            if p.startswith("http"):
                current["url"] = p
            else:
                current["note"] = clean(f"{current['note']} {p}")
            continue
        # Generic prose/list formats: one record per non-heading paragraph.
        if mode == "generic":
            seq["generic"] += 1
            records.append(_record_from_parts(rid=f"M{seq['generic']}", citation=p, section=section, mode="current", source_file=source_file))
    flush()
    return records


def _parse_docx(path: Path) -> list[dict[str, Any]]:
    if Document is None:
        raise RuntimeError("DOCX ingestion requires python-docx")
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    # Tables are common in manually curated bibliographies; append each row as a delimited record.
    for table in doc.tables:
        for row in table.rows:
            vals = [clean(c.text) for c in row.cells]
            if any(vals):
                paragraphs.append(" | ".join(vals))
    return _parse_paragraph_records(paragraphs, path.name)


def _parse_pdf(path: Path) -> list[dict[str, Any]]:
    reader = PdfReader(str(path))
    paras: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        paras.extend(x for x in text.splitlines() if clean(x))
    return _parse_paragraph_records(paras, path.name)


def _row_to_record(row: dict[str, Any], idx: int, source_file: str) -> dict[str, Any]:
    lower = {str(k).lower().strip(): v for k, v in row.items()}
    rid = clean(lower.get("id") or lower.get("manual_id") or f"M{idx}")
    title = clean(lower.get("title") or lower.get("name"))
    citation = clean(lower.get("citation") or lower.get("reference") or title)
    url = clean(lower.get("url") or lower.get("link") or lower.get("doi_url"))
    note = clean(lower.get("note") or lower.get("notes") or lower.get("comment"))
    rec = _record_from_parts(rid=rid, citation=citation, url=url, note=note, section=clean(lower.get("section")), mode="current", source_file=source_file)
    for key in ("authors", "source", "date", "doi"):
        if clean(lower.get(key)):
            rec[key] = clean(lower[key])
    if clean(lower.get("date")):
        raw_date = clean(lower.get("date"))
        if re.fullmatch(r"20\d{2}-\d{2}-\d{2}", raw_date):
            rec["date_precision"] = "day"
        elif re.fullmatch(r"20\d{2}-\d{2}", raw_date):
            rec["date_precision"] = "month"
            rec["date"] = raw_date + "-01"
        elif re.fullmatch(r"20\d{2}", raw_date):
            rec["date_precision"] = "year"
            rec["date"] = raw_date + "-01-01"
        else:
            rec["date_precision"] = _date_precision(raw_date) or rec.get("date_precision", "")
    if title:
        rec["title"] = title
    status = _norm(lower.get("status"))
    if "forthcoming" in status or "unpublished" in status:
        rec["manual_record_status"] = "forthcoming_unpublished"
    return rec


def parse_manual_file(path: str | Path) -> list[dict[str, Any]]:
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"Unsupported manual-ingest format: {ext or '(none)'}")
    if ext == ".docx":
        return _parse_docx(p)
    if ext == ".pdf":
        return _parse_pdf(p)
    if ext == ".csv":
        with p.open("r", encoding="utf-8-sig", newline="") as f:
            return [_row_to_record(row, i, p.name) for i, row in enumerate(csv.DictReader(f), 1)]
    if ext in {".json", ".yaml", ".yml"}:
        if ext == ".json":
            data = json.loads(p.read_text(encoding="utf-8"))
        else:
            if yaml is None:
                raise RuntimeError("YAML ingestion requires PyYAML")
            data = yaml.safe_load(p.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            data = data.get("items") or data.get("papers") or data.get("records") or [data]
        if not isinstance(data, list):
            raise ValueError("JSON/YAML manual ingest must contain a list or an items/papers/records list")
        out = []
        for i, row in enumerate(data, 1):
            out.append(_row_to_record(row if isinstance(row, dict) else {"citation": str(row)}, i, p.name))
        return out
    return _parse_paragraph_records(p.read_text(encoding="utf-8-sig").splitlines(), p.name)


def _corpus_entries(state: dict[str, Any]) -> list[tuple[str, dict[str, Any]]]:
    out = []
    for strand in ("strand_a", "strand_b", "strand_c", "frontier_evidence"):
        for item in state.get(strand, []) if isinstance(state.get(strand), list) else []:
            if isinstance(item, dict):
                out.append((strand, item))
    return out


def _title_similarity(a: str, b: str) -> float:
    # Deterministic token overlap + prefix containment; deliberately conservative.
    na, nb = sr.norm_title(a), sr.norm_title(b)
    if not na or not nb:
        return 0.0
    if na == nb:
        return 1.0
    if na in nb or nb in na:
        return 0.96
    sa, sb = set(na.split()), set(nb.split())
    j = len(sa & sb) / max(1, len(sa | sb))
    contain = len(sa & sb) / max(1, min(len(sa), len(sb)))
    return max(j, contain * 0.88)


def match_existing(record: dict[str, Any], state: dict[str, Any]) -> tuple[str, dict[str, Any] | None, float]:
    rdoi = clean(record.get("doi")).lower()
    rurl = _canonical_url(record.get("url", ""))
    best: tuple[str, dict[str, Any] | None, float] = ("", None, 0.0)
    for strand, item in _corpus_entries(state):
        idoi = _doi(clean(item.get("_doi") or item.get("link", "")))
        if rdoi and idoi and rdoi == idoi:
            return strand, item, 1.0
        iurl = _canonical_url(item.get("link", ""))
        if rurl and iurl and rurl == iurl:
            return strand, item, 1.0
        sim = _title_similarity(record.get("title", ""), item.get("title", ""))
        if sim > best[2]:
            best = (strand, item, sim)
    # Title-only fallback is deliberately strict. Different sovereignty/strategy
    # papers often share most topic words and must not collapse into one record.
    return best if best[2] >= 0.94 else ("", None, best[2])


def _seen_automatically(record: dict[str, Any], state: dict[str, Any]) -> bool:
    rurl = _canonical_url(record.get("url", ""))
    if not rurl:
        return False
    seen = state.get("scan_state", {}).get("institution_seen_fingerprints", {})
    if not isinstance(seen, dict):
        return False
    for key in seen:
        if _canonical_url(str(key).split("|", 1)[0]) == rurl:
            return True
    return False


def _source_profile(record: dict[str, Any]) -> tuple[str, int, str, str]:
    url = clean(record.get("url"))
    domain = urlparse(url).netloc.lower().removeprefix("www.") if url else ""
    configured = sr.institution_source_for_domain(domain) if domain else None
    if configured:
        name, tier = configured
        return clean(record.get("source") or name), int(tier), f"Tier {tier}", "institutional"
    note = _norm(f"{record.get('citation','')} {record.get('curator_note','')}")
    if domain.endswith("europa.eu") or domain in {"consilium.europa.eu", "data.europa.eu", "era.gv.at"}:
        return clean(record.get("source") or domain), 1, "Tier 1", "institutional"
    if record.get("doi") or any(x in note for x in ("academic", "journal", "cambridge university press", "peer-reviewed")):
        return clean(record.get("source") or domain), 2, "Tier 2 comparable", "scholarly"
    if "type 1" in note:
        return clean(record.get("source") or domain), 1, "Tier 1", "institutional"
    if "type 2" in note or "type 3" in note:
        return clean(record.get("source") or domain), 2, "Tier 2 comparable", "institutional"
    if any(x in domain for x in ("industrialinfo.com", "jonesday.com", "noerr.com", "insideglobaltech.com", "zabala.eu")):
        return clean(record.get("source") or domain), 3, "Tier 3 specialist", "institutional"
    return clean(record.get("source") or domain), 2, "Tier 2 comparable", "institutional"


def _pdf_bytes_text(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception:
        return ""
    out = []
    for page in reader.pages[:70]:
        try:
            out.append(page.extract_text() or "")
        except Exception:
            pass
    return clean(" ".join(out))


def retrieve_source(record: dict[str, Any], *, timeout: int = 18, session: requests.Session | None = None) -> dict[str, Any]:
    """Retrieve a directly cited source.  No search engine is used here.

    A generic homepage or a page that cannot be tied to the candidate title stays
    uncertain/metadata-only.  This prevents the curator's note from becoming evidence.
    """
    url = clean(record.get("url"))
    if not url:
        return {"retrieval_status": "no_url", "evidence_status": "uncertain_record", "text_mode": "metadata_only", "abstract": "", "body": "", "resolved_url": ""}
    s = session or sr.SESSION
    try:
        r = s.get(url, timeout=timeout, allow_redirects=True)
    except Exception as e:
        return {"retrieval_status": f"request_failed:{type(e).__name__}", "evidence_status": "uncertain_record", "text_mode": "metadata_only", "abstract": "", "body": "", "resolved_url": url}
    if r.status_code != 200:
        return {"retrieval_status": f"http_{r.status_code}", "evidence_status": "uncertain_record", "text_mode": "metadata_only", "abstract": "", "body": "", "resolved_url": clean(r.url or url)}
    ctype = _norm(r.headers.get("content-type", ""))
    resolved = clean(r.url or url)
    title = ""; abstract = ""; body = ""; published = ""; authors = ""; pdf_url = ""
    if "pdf" in ctype or resolved.lower().split("?", 1)[0].endswith(".pdf"):
        body = _pdf_bytes_text(r.content)
        title = clean(record.get("title"))
    else:
        soup = BeautifulSoup(r.text, "html.parser")
        title = sr.meta_content(soup, ["og:title", "twitter:title", "citation_title", "dc.title"]) or clean(soup.h1.get_text(" ", strip=True) if soup.h1 else "") or clean(soup.title.get_text(" ", strip=True) if soup.title else "")
        abstract = sr.meta_content(soup, ["description", "og:description", "twitter:description", "citation_abstract", "dc.description"])
        published = sr.meta_content(soup, ["article:published_time", "datePublished", "citation_publication_date", "DC.date", "date"])
        authors = sr.meta_content(soup, ["citation_author", "author", "dc.creator"])
        article_body = ""
        for script in soup.find_all("script", attrs={"type": re.compile("ld\\+json", re.I)}):
            try:
                data = json.loads(script.string or script.get_text())
            except Exception:
                continue
            for obj in sr.jsonld_objects(data):
                if not article_body and obj.get("articleBody"):
                    article_body = clean(obj.get("articleBody"))
                if not published and obj.get("datePublished"):
                    published = clean(obj.get("datePublished"))
        for a in soup.find_all("a", href=True):
            href = urljoin(resolved, a.get("href"))
            label = _norm(a.get_text(" ", strip=True))
            if href.lower().split("?", 1)[0].endswith(".pdf") or label in {"pdf", "download pdf", "download report", "download paper"}:
                pdf_url = href; break
        for bad in soup(["script", "style", "nav", "header", "footer", "aside", "form", "noscript"]):
            bad.decompose()
        container = soup.find("article") or soup.find("main") or soup.body
        body = article_body or clean(container.get_text(" ", strip=True) if container else "")
        # Prefer linked full PDF where the landing page is only metadata/abstract.
        if pdf_url and len(body.split()) < 1800:
            try:
                pr = s.get(pdf_url, timeout=timeout, allow_redirects=True)
                if pr.status_code == 200 and len(pr.content) <= 30_000_000:
                    pbody = _pdf_bytes_text(pr.content)
                    if len(pbody.split()) > len(body.split()):
                        body = pbody
                        resolved = clean(pr.url or pdf_url)
            except Exception:
                pass
    words = len(body.split())
    title_sim = _title_similarity(record.get("title", ""), title)
    # For PDF links the record title is the only title we may have; direct PDF retrieval is still primary evidence.
    direct_pdf = "pdf" in ctype or resolved.lower().split("?", 1)[0].endswith(".pdf")
    path = urlparse(clean(record.get("url"))).path.strip("/")
    generic_homepage = not path and not direct_pdf
    verified = bool((direct_pdf and words >= 120) or (title_sim >= 0.58 and (words >= 120 or len(abstract.split()) >= 35))) and not generic_homepage
    if record.get("manual_secondary_hint"):
        evidence_status = "secondary_reference"
    elif verified:
        evidence_status = "verified_primary_source"
    elif words or abstract:
        evidence_status = "uncertain_record"
    else:
        evidence_status = "uncertain_record"
    if words >= 650:
        text_mode = "full_text"
    elif words >= 80:
        text_mode = "partial_text"
    elif len(abstract.split()) >= 35:
        text_mode = "abstract_only"
    else:
        text_mode = "metadata_only"
    return {
        "retrieval_status": "retrieved",
        "evidence_status": evidence_status,
        "text_mode": text_mode,
        "title": title,
        "authors": authors,
        "published": published,
        "abstract": abstract,
        "body": body,
        "resolved_url": resolved,
        "title_similarity": round(title_sim, 3),
        "word_count": words,
    }


def load_review_evidence(path: str | Path | None) -> dict[str, dict[str, Any]]:
    """Load a transparent, machine-readable source-review pack.

    The pack is not a manual-admission override.  It records evidence obtained from
    the underlying source by a reviewer when the repository runtime cannot retrieve
    that source itself.  The normal substantive gate is still run on the reviewed
    source extract/abstract.  Matrix annotations are accepted only when the review
    explicitly marks the classification as source-evidence-based.
    """
    if not path:
        return {}
    p = Path(path)
    raw = json.loads(p.read_text(encoding="utf-8"))
    items = raw.get("items", raw) if isinstance(raw, dict) else {}
    if not isinstance(items, dict):
        raise ValueError("review evidence must be an object keyed by manual_id or contain an 'items' object")
    out: dict[str, dict[str, Any]] = {}
    for key, value in items.items():
        if isinstance(value, dict):
            out[clean(key)] = deepcopy(value)
    return out


def _review_hash(review_evidence: dict[str, dict[str, Any]] | None) -> str:
    if not review_evidence:
        return ""
    payload = json.dumps(review_evidence, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _review_url_matches(record: dict[str, Any], review: dict[str, Any] | None) -> bool:
    """Bind cached review evidence to the exact URL supplied by the curator.

    Manual review packs are repository-controlled evidence caches, not a discovery
    mechanism.  A review may only affect a record when it names the same supplied
    URL.  This prevents a review for a similarly titled paper from silently being
    reused and keeps the manual lane search-engine-free.
    """
    if not review:
        return False
    supplied = _canonical_url(record.get("url", ""))
    reviewed = _canonical_url(review.get("review_source_url") or review.get("source_url") or "")
    return bool(supplied and reviewed and supplied == reviewed)


def _review_core_gate_verified(record: dict[str, Any], review: dict[str, Any] | None, retrieval: dict[str, Any]) -> bool:
    """Return True only for an explicitly adjudicated underlying-source pass.

    This is an alternate *substantive* evaluator, not a curator-cell shortcut.  It
    exists because the lexical scanner gate is a high-precision heuristic and must
    not become a mandatory keyword gate after a human/LLM review of the underlying
    source has established EU/European R&I in geopolitical context.
    """
    if not _review_url_matches(record, review):
        return False
    if not review or not review.get("core_gate_verified"):
        return False
    if retrieval.get("evidence_status") != "verified_primary_source":
        return False
    status = clean(review.get("review_status"))
    if not status.startswith("reviewed_pass_core_gate"):
        return False
    evidence_text = clean(review.get("evidence_text") or review.get("abstract") or review.get("body"))
    return len(evidence_text.split()) >= 25


def _apply_review_core_gate(ev: dict[str, Any], review: dict[str, Any], retrieval: dict[str, Any]) -> dict[str, Any]:
    """Overlay a transparent reviewed-source pass while retaining scanner diagnostics."""
    out = deepcopy(ev)
    out["scanner_a_pass"] = bool(ev.get("a_pass"))
    out["scanner_aboutness_reason"] = clean(ev.get("aboutness_reason"))
    out.update({
        "a_pass": True,
        "a_focus_pass": True,
        "aboutness_pass": True,
        "aboutness_reason": "reviewed_underlying_source_substantive_gate",
        "eu_relevance": "direct",
        "eu_evidence": ["reviewed underlying source establishes EU/European scope"],
        "ri_evidence": ["reviewed underlying source establishes substantive R&I mechanism"],
        "geo_evidence": ["reviewed underlying source establishes geopolitical/economic-security mechanism"],
        "bridge_sentence": clean(review.get("core_gate_basis") or review.get("review_basis") or review.get("evidence_text"))[:420],
        "a_route": "reviewed_underlying_source",
        "bridge_supported": True,
        "bridge_mode": "reviewed_source",
    })
    return out


def _review_core_gate_failed(record: dict[str, Any], review: dict[str, Any] | None) -> bool:
    """Recognise an explicit source-based A-gate rejection bound to the supplied URL."""
    if not _review_url_matches(record, review) or not review:
        return False
    return review.get("core_gate_verified") is False and clean(review.get("review_status")).startswith("reviewed_fail_core_gate")


def _apply_review_core_gate_fail(ev: dict[str, Any], review: dict[str, Any]) -> dict[str, Any]:
    """Prevent a lexical false positive from overriding reviewed underlying-source evidence."""
    out = deepcopy(ev)
    out["scanner_a_pass"] = bool(ev.get("a_pass"))
    out["scanner_aboutness_reason"] = clean(ev.get("aboutness_reason"))
    out.update({
        "a_pass": False,
        "a_focus_pass": False,
        "aboutness_pass": False,
        "aboutness_reason": "reviewed_underlying_source_failed_substantive_gate",
        "a_route": "",
        "bridge_supported": False,
        "bridge_mode": "",
        "reviewed_gate_basis": clean(review.get("review_basis") or review.get("evidence_text"))[:420],
    })
    return out


def _record_with_review(record: dict[str, Any], review: dict[str, Any] | None) -> dict[str, Any]:
    """Apply verified bibliographic corrections without erasing curator provenance."""
    out = deepcopy(record)
    if not review or not _review_url_matches(record, review):
        return out
    overrides = {
        "title": "title", "authors": "authors", "source": "source",
        "doi": "doi", "date": "published", "date_precision": "date_precision",
        "manual_record_status": "record_status",
    }
    for dest, src in overrides.items():
        value = review.get(src)
        if value not in {None, ""}:
            out[dest] = clean(value) if isinstance(value, str) else value
    if review.get("published") and not review.get("date_precision"):
        parsed = sr.parse_date(review.get("published"))
        if parsed:
            out["date_precision"] = "day"
            out["date"] = parsed.isoformat()
    if review.get("resolved_url"):
        out["review_resolved_url"] = clean(review.get("resolved_url"))
    if review.get("resolved_primary") or (review.get("source_verified") and review.get("primary_source")):
        # Preserve the original flag in diagnostics, but evaluate the resolved primary source.
        out["manual_secondary_hint"] = False
    if review.get("source_verified") and sr.parse_date(review.get("published")):
        out["manual_verification_required"] = False
    return out


def _retrieval_from_review(record: dict[str, Any], review: dict[str, Any] | None) -> dict[str, Any] | None:
    if not review or not _review_url_matches(record, review):
        return None
    abstract = clean(review.get("abstract") or review.get("evidence_text"))
    body = clean(review.get("body"))
    primary = bool(review.get("primary_source"))
    verified = bool(review.get("source_verified"))
    corroborated_event = bool(review.get("corroborated_current_event"))
    if verified and primary:
        evidence_status = "verified_primary_source"
    elif verified and corroborated_event:
        evidence_status = "verified_corroborated_current_event"
    elif review.get("secondary_reference"):
        evidence_status = "secondary_reference"
    else:
        evidence_status = "uncertain_record"
    text_mode = clean(review.get("text_mode"))
    if text_mode not in {"full_text", "partial_text", "abstract_only", "metadata_only"}:
        text_mode = "full_text" if len(body.split()) >= 650 else ("partial_text" if len(body.split()) >= 80 else ("abstract_only" if len(abstract.split()) >= 35 else "metadata_only"))
    return {
        "retrieval_status": clean(review.get("retrieval_status")) or "reviewed_exact_supplied_url_cache",
        "evidence_status": evidence_status,
        "text_mode": text_mode,
        "title": clean(review.get("title") or record.get("title")),
        "authors": clean(review.get("authors") or record.get("authors")),
        "published": clean(review.get("published") or record.get("date")),
        "abstract": abstract,
        "body": body,
        "resolved_url": clean(record.get("url")),
        "review_resolved_url": clean(review.get("resolved_url")),
        "review_source_url": clean(review.get("review_source_url") or review.get("source_url")),
        "title_similarity": 1.0 if clean(review.get("title") or record.get("title")) else 0.0,
        "word_count": len(f"{abstract} {body}".split()),
        "review_status": clean(review.get("review_status")),
        "review_basis": clean(review.get("review_basis")),
        "reviewed_at": clean(review.get("reviewed_at")),
    }


def _matrix_fields_from_review(record: dict[str, Any], review: dict[str, Any] | None) -> dict[str, Any]:
    if not review or not _review_url_matches(record, review) or not review.get("matrix_evidence_verified"):
        return {}
    dim = clean(review.get("matrix_dimension")).lower()
    implied = clean(review.get("quadrant_implied")).upper()
    claimed = clean(review.get("quadrant_claimed")).upper()
    if dim not in MATRIX_DIMENSIONS or implied not in MATRIX_QUADRANTS:
        return {}
    if claimed and claimed not in MATRIX_QUADRANTS:
        claimed = ""
    curator = clean(record.get("curator_primary_cell")).upper()
    curator_dim = MATRIX_ROWS.get(curator[:1], "") if len(curator) == 3 and curator[1:2] == "-" else ""
    curator_quad = curator[-1:] if curator_dim else ""
    agreement = ""
    if curator:
        agreement = "agrees" if curator_dim == dim and curator_quad == implied else "differs"
    fields = {
        "matrix_dimension": dim,
        "quadrant_implied": implied,
        "matrix_quadrant": implied,
        "matrix_classification_source": "reviewed_underlying_source",
        "matrix_evidence_basis": clean(review.get("matrix_basis") or review.get("display_claim")),
        "curator_primary_cell": curator,
        "curator_cells": list(record.get("curator_cells") or []),
        "curator_cell_agreement": agreement,
    }
    if claimed:
        fields["quadrant_claimed"] = claimed
    return fields


def _gate_reason(ev: dict[str, Any]) -> str:
    if ev.get("a_pass") or ev.get("b_pass"):
        return "passes_substantive_gate"
    reason = clean(ev.get("aboutness_reason"))
    if reason:
        return reason
    if ev.get("eu_relevance") != "direct" and not ev.get("b_pass"):
        return "no_direct_eu"
    if not ev.get("ri_evidence"):
        return "no_substantive_ri"
    if not ev.get("geo_evidence"):
        return "no_geopolitical_context"
    return "substantive_gate_failed"


def _provenance_value(item: dict[str, Any], manual_id: str) -> None:
    existing_label = _norm(item.get("discovery_provenance"))
    existing = item.get("provenance") or []
    if isinstance(existing, str):
        existing = [existing]
    existing_norm = {_norm(x) for x in existing if clean(x)}
    manual_only = existing_label == "manual" or ("manual_candidate_ingestion" in existing_norm and "automated_discovery" not in existing_norm)
    if manual_only:
        item["provenance"] = ["manual_candidate_ingestion"]
        item["discovery_provenance"] = "manual"
    else:
        # Legacy corpus items without provenance came from automated discovery.
        item["provenance"] = ["automated_discovery", "manual_candidate_ingestion"]
        item["discovery_provenance"] = "both"
    mids = item.get("manual_ingest_ids") if isinstance(item.get("manual_ingest_ids"), list) else []
    if manual_id not in mids:
        mids.append(manual_id)
    item["manual_ingest_ids"] = mids


def _effective_publication_date(record: dict[str, Any], retrieval: dict[str, Any]) -> dt.date | None:
    """Return a publication date suitable for admission, without inventing precision.

    A retrieved source date takes precedence.  Otherwise day/month precision from the
    curator can be used unless the curator explicitly marked the date/source as needing
    verification.  A bare year is never enough for a bounded-window admission.
    """
    fetched = sr.parse_date(retrieval.get("published"))
    if fetched:
        return fetched
    if record.get("manual_verification_required"):
        return None
    if record.get("date_precision") not in {"day", "month"}:
        return None
    return sr.parse_date(record.get("date"))


def _new_public_item(record: dict[str, Any], retrieval: dict[str, Any], ev: dict[str, Any], strand: str, ingested_at: str, *, publication_date: dt.date | None = None, review: dict[str, Any] | None = None) -> dict[str, Any]:
    source_name, tier, tier_label, source_kind = _source_profile(record)
    date = publication_date or _effective_publication_date(record, retrieval)
    if not date:
        raise ValueError("verified manual source has no usable publication date")
    text = clean(f"{record.get('title','')}. {retrieval.get('abstract','')}. {retrieval.get('body','')[:45000]}")
    item = sr.build_item(
        title=record.get("title") or retrieval.get("title") or "Untitled manual candidate",
        authors=record.get("authors") or retrieval.get("authors") or source_name,
        source=source_name or urlparse(record.get("url", "")).netloc,
        date=date,
        link=(clean(retrieval.get("review_resolved_url")) if review and review.get("resolved_primary") else "") or record.get("url") or retrieval.get("resolved_url") or (f"https://doi.org/{record.get('doi')}" if record.get("doi") else ""),
        item_type="manual-verified scholarly/policy source",
        strand=strand,
        evidence=ev,
        source_rank=float(tier), tier_label=tier_label,
        text=text, doi=record.get("doi", ""), preprint=False,
    )
    pub = sr.public_item(item, new_this_scan=False, first_seen=ingested_at)
    pub.update({
        "provenance": ["manual_candidate_ingestion"],
        "discovery_provenance": "manual",
        "manual_ingest_ids": [record.get("manual_id")],
        "manual_ingested_at": ingested_at,
        "evidence_status": retrieval.get("evidence_status"),
        "source_text_mode": retrieval.get("text_mode"),
        "manual_supplied_url": clean(record.get("url")),
        "review_resolved_url": clean(retrieval.get("review_resolved_url")),
    })
    if review:
        pub.update({
            "source_review_status": clean(review.get("review_status") or "reviewed"),
            "source_reviewed_at": clean(review.get("reviewed_at")),
            "source_review_basis": clean(review.get("review_basis")),
        })
        claim = clean(review.get("display_claim"))
        if claim:
            reviewed_summary = clean(review.get("summary") or pub.get("summary") or claim)
            pub["core_message"] = sr.plain_language_claim(reviewed_summary, pub.get("title", ""), claim)
            pub["summary"] = reviewed_summary
        pub.update(_matrix_fields_from_review(record, review))
    return pub


def _new_public_signal(record: dict[str, Any], retrieval: dict[str, Any], state: dict[str, Any], ingested_at: str, *, publication_date: dt.date, review: dict[str, Any] | None = None) -> dict[str, Any] | None:
    """Apply the normal factual/anchoring checks before admitting a manual weak signal."""
    source_name, _, _, _ = _source_profile(record)
    desc = clean(f"{retrieval.get('abstract','')} {retrieval.get('body','')[:18000]}")
    title = clean(record.get("title") or retrieval.get("title"))
    if not title or not sr.factual_news(title, desc):
        return None
    full = clean(f"{title} {desc}")
    candidate = {
        "headline": title,
        "source": source_name or urlparse(record.get("url", "")).netloc,
        "date": publication_date.isoformat(),
        "link": clean((retrieval.get("review_resolved_url") if review and review.get("resolved_primary") else "") or record.get("url") or retrieval.get("resolved_url")),
        "language": "en",
        "_desc": desc,
        "_themes": sr.themes_for(full),
        "_entities": sr.distinct_matches(full, sr.ENTITY_TERMS + sr.GEO_ACTORS),
    }
    anchored = sr.anchor_news([candidate], state.get("strand_a", []))
    if not anchored:
        return None
    item = anchored[0]
    item.update({
        "first_seen": ingested_at,
        "new_this_scan": False,
        "core_message": sr.plain_language_claim(desc, title, clean(item.get("what") or title)),
        "provenance": ["manual_candidate_ingestion"],
        "discovery_provenance": "manual",
        "manual_ingest_ids": [record.get("manual_id")],
        "manual_ingested_at": ingested_at,
        "evidence_status": retrieval.get("evidence_status"),
        "source_text_mode": retrieval.get("text_mode"),
        "manual_supplied_url": clean(record.get("url")),
        "review_resolved_url": clean(retrieval.get("review_resolved_url")),
    })
    if review:
        claim = clean(review.get("display_claim"))
        if claim:
            simple_claim = sr.plain_language_claim(desc, title, claim)
            item["core_message"] = simple_claim
            item["what"] = simple_claim
            item["signal_note"] = clean(f"{simple_claim.rstrip('. ')}. {item.get('why_it_matters','')}")
        item.update({
            "source_review_status": clean(review.get("review_status") or "reviewed"),
            "source_reviewed_at": clean(review.get("reviewed_at")),
            "source_review_basis": clean(review.get("review_basis")),
        })
        item.update(_matrix_fields_from_review(record, review))
    return item


def _append_signal_if_new(state: dict[str, Any], item: dict[str, Any]) -> bool:
    arr = state.setdefault("strand_c", [])
    ident = sr.signal_identity(item)
    for old in arr:
        if sr.signal_identity(old) == ident:
            return False
    arr.append(item)
    return True


def _append_if_new(state: dict[str, Any], strand_key: str, item: dict[str, Any]) -> bool:
    arr = state.setdefault(strand_key, [])
    ident = sr.stable_item_identity(item.get("title", ""), item.get("link", ""))
    for old in arr:
        if sr.stable_item_identity(old.get("title", ""), old.get("link", "")) == ident:
            return False
    arr.append(item)
    return True


def _saved_miss_reason(record: dict[str, Any], state: dict[str, Any], automated_status: str, retrieval: dict[str, Any], ev: dict[str, Any]) -> str:
    """Explain what the saved state can actually establish about a manual candidate miss."""
    if automated_status == "found_in_corpus":
        return "not_missed"
    record_status = clean(record.get("manual_record_status"))
    if record_status == "forthcoming_unpublished":
        return "forthcoming_not_scanner_recall_target"
    if record_status == "context_outside_primary_window":
        return "context_only_not_scanner_recall_target"
    evidence_status = retrieval.get("evidence_status")
    retrieval_status = retrieval.get("retrieval_status")
    if evidence_status == "secondary_reference":
        return "secondary_reference_requires_primary_verification"
    if evidence_status == "verified_primary_source":
        if ev.get("a_pass") or ev.get("b_pass"):
            return "discovery_recall_gap"
        return f"recovered_source_failed_substantive_gate:{_gate_reason(ev)}"
    if retrieval_status not in {"not_attempted_offline", "not_needed_existing_corpus", ""} and retrieval.get("text_mode") == "metadata_only":
        return "underlying_source_retrieval_failed_or_insufficient"
    if automated_status == "scanner_seen_url_not_admitted":
        # The supplied state persists the seen fingerprint but not a per-URL rejection trace.
        return "scanner_saw_url_but_saved_state_has_no_item_level_rejection_reason"

    raw = clean(record.get("url"))
    host = urlparse(raw if "://" in raw else "https://" + raw).netloc.lower().removeprefix("www.") if raw else ""
    warnings = state.get("scan_diagnostics", {}).get("source_warnings", []) if isinstance(state.get("scan_diagnostics"), dict) else []
    warning_domains = set()
    for warning in warnings if isinstance(warnings, list) else []:
        m = re.search(r"No usable sitemap:\s*([^\s]+)", clean(warning), re.I)
        if m:
            warning_domains.add(m.group(1).lower().removeprefix("www."))
    if host and any(host == d or host.endswith("." + d) or d.endswith("." + host) for d in warning_domains):
        return "direct_source_sitemap_unavailable_in_saved_scan"

    seen_fingerprints = state.get("scan_state", {}).get("institution_seen_fingerprints", {}) if isinstance(state.get("scan_state"), dict) else {}
    seen_domains = set()
    if isinstance(seen_fingerprints, dict):
        for fingerprint in seen_fingerprints:
            raw_seen = clean(fingerprint).split("|", 1)[0]
            try:
                seen_host = urlparse(raw_seen).netloc.lower().removeprefix("www.")
            except Exception:
                seen_host = ""
            if seen_host:
                seen_domains.add(seen_host)
    if host and any(host == d or host.endswith("." + d) or d.endswith("." + host) for d in seen_domains):
        return "not_observed_during_bounded_direct_source_rotation"
    if record.get("doi"):
        return "not_observed_in_saved_scholarly_index_scan"
    if host:
        return "source_not_observed_in_saved_direct_discovery_ledger"
    return "insufficient_bibliographic_locator_for_discovery_diagnosis"



def _recall_failure_category(record: dict[str, Any], state: dict[str, Any], automated_status: str) -> str:
    """Classify the likely discovery-stage failure using only persisted scanner evidence.

    This deliberately does not infer pass-1 or quality rejection reasons that the saved
    state did not retain. It is a recall diagnostic, not a substantive admission gate.
    """
    if automated_status == "found_in_corpus":
        return "not_missed"
    if automated_status == "scanner_seen_url_not_admitted":
        return "seen_but_rejection_stage_unknown_from_saved_state"
    raw = clean(record.get("url"))
    host = urlparse(raw if "://" in raw else "https://" + raw).netloc.lower().removeprefix("www.") if raw else ""
    if record.get("doi"):
        return "scholarly_index_item_not_observed"
    added_domains = {clean(x).lower().removeprefix("www.") for x in sr.CONFIG.get("manual_recall_added_domains", [])}
    if host and any(host == d or host.endswith("." + d) for d in added_domains):
        return "source_not_covered_prior_to_targeted_manual_recall_expansion"
    configured = sr.institution_source_for_domain(host) if host else None
    if configured:
        warnings = state.get("scan_diagnostics", {}).get("source_warnings", []) if isinstance(state.get("scan_diagnostics"), dict) else []
        for warning in warnings if isinstance(warnings, list) else []:
            m = re.search(r"No usable sitemap:\s*([^\s]+)", clean(warning), re.I)
            if m:
                d = m.group(1).lower().removeprefix("www.")
                if host == d or host.endswith("." + d) or d.endswith("." + host):
                    return "sitemap_or_feed_failure"
        seen = state.get("scan_state", {}).get("institution_seen_fingerprints", {}) if isinstance(state.get("scan_state"), dict) else {}
        if isinstance(seen, dict):
            for fingerprint in seen:
                seen_url = clean(fingerprint).split("|", 1)[0]
                seen_host = urlparse(seen_url).netloc.lower().removeprefix("www.") if seen_url else ""
                if seen_host and (host == seen_host or host.endswith("." + seen_host) or seen_host.endswith("." + host)):
                    return "covered_source_exact_item_not_observed_in_rotation"
        return "covered_source_but_saved_state_cannot_localise_item_miss"
    if host:
        return "source_not_covered_by_direct_discovery"
    return "unknown_insufficient_locator"

def _diagnostic_record(record: dict[str, Any], state: dict[str, Any], automated_status: str, retrieval: dict[str, Any], ev: dict[str, Any], decision: str, matched: dict[str, Any] | None, *, links_validated: bool = False, review: dict[str, Any] | None = None, original_record: dict[str, Any] | None = None) -> dict[str, Any]:
    reason = _gate_reason(ev)
    miss = _saved_miss_reason(record, state, automated_status, retrieval, ev)
    out = {
        "manual_id": record.get("manual_id"),
        "title": record.get("title"),
        "date": record.get("date"),
        "url": record.get("url"),
        "section": record.get("section"),
        "manual_record_status": record.get("manual_record_status"),
        "manual_candidate_kind": record.get("manual_candidate_kind", "substantive"),
        "manual_verification_required": bool(record.get("manual_verification_required")),
        "manual_link_status": "user_validated_reachable" if links_validated and record.get("url") else ("supplied_unchecked" if record.get("url") else "not_supplied"),
        "curator_cells": record.get("curator_cells", []),
        "curator_primary_cell": record.get("curator_primary_cell", ""),
        "curator_cell_mapping_status": record.get("curator_cell_mapping_status", ""),
        "automated_status": automated_status,
        "evidence_status": retrieval.get("evidence_status"),
        "retrieval_status": retrieval.get("retrieval_status"),
        "text_mode": retrieval.get("text_mode"),
        # Keep both the curator-supplied URL and any primary/full-text URL reached
        # through the permitted exact-link chain.  The supplied URL stays in ``url``;
        # a resolved primary is additional provenance, never a silent replacement.
        "resolved_url": clean(retrieval.get("resolved_url")),
        "review_resolved_url": clean(retrieval.get("review_resolved_url")),
        "gate": {
            "a_pass": bool(ev.get("a_pass")), "b_pass": bool(ev.get("b_pass")),
            "eu_relevance": ev.get("eu_relevance"), "aboutness_reason": ev.get("aboutness_reason"),
            "reason": reason,
        },
        "decision": decision,
        "miss_reason": miss,
        "recall_failure_category": _recall_failure_category(record, state, automated_status),
        "matched_title": clean(matched.get("title")) if matched else "",
    }
    if original_record:
        corrections = {}
        for key in ("title", "date", "url", "authors", "source", "doi", "manual_record_status"):
            before = original_record.get(key)
            after = record.get(key)
            if clean(before) != clean(after):
                corrections[key] = {"supplied": before, "reviewed": after}
        if corrections:
            out["reviewed_bibliographic_corrections"] = corrections
    if review:
        out.update({
            "source_review_status": clean(review.get("review_status")),
            "source_reviewed_at": clean(review.get("reviewed_at")),
            "source_review_basis": clean(review.get("review_basis")),
            "review_url_bound_to_supplied_link": _review_url_matches(original_record or record, review),
            "core_gate_verified": bool(review.get("core_gate_verified")) and _review_url_matches(original_record or record, review),
            "matrix_evidence_verified": bool(review.get("matrix_evidence_verified")) and _review_url_matches(original_record or record, review),
            "matrix_review": _matrix_fields_from_review(original_record or record, review),
            "direct_links_followed": list(review.get("direct_links_followed") or []),
        })
    return out


def _merge_record_history(old: list[dict[str, Any]], new: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_key = {}
    for x in old + new:
        if not isinstance(x, dict):
            continue
        key = f"{clean(x.get('source_file'))}|{clean(x.get('manual_id'))}|{sr.norm_title(x.get('title',''))}"
        by_key[key] = x
    return list(by_key.values())


def apply_manual_ingest(state: dict[str, Any], records: list[dict[str, Any]], *, source_path: str | Path, fetch: bool = True, refresh: bool = False, links_validated: bool = False, review_evidence: dict[str, dict[str, Any]] | None = None, now: dt.datetime | None = None, session: requests.Session | None = None) -> tuple[dict[str, Any], dict[str, Any]]:
    """Apply manual candidates to a copy of ``state`` while preserving scan timestamps."""
    out = deepcopy(state)
    # Keep the same EU-context anchor set as the automatic scanner so the exceptional
    # external-shock route behaves identically for reviewed manual candidates.
    sr.ACTIVE_EU_CONTEXT_ANCHORS = [dict(x) for x in out.get("strand_a", []) if isinstance(x, dict)]
    old_last_updated = out.get("last_updated")
    now = now or dt.datetime.now(dt.timezone.utc)
    if now.tzinfo is None:
        now = now.replace(tzinfo=dt.timezone.utc)
    ingested_at = now.astimezone(dt.timezone.utc).isoformat(timespec="minutes").replace("+00:00", "Z")
    source = Path(source_path)
    sha = _file_hash(source)
    review_evidence = review_evidence or {}
    review_sha = _review_hash(review_evidence)
    batch_id = sha[:12] + ("-" + review_sha[:8] if review_sha else "")
    manual = out.get("manual_ingest") if isinstance(out.get("manual_ingest"), dict) else {}
    batches = manual.get("batches") if isinstance(manual.get("batches"), list) else []
    if not refresh and any(b.get("sha256") == sha and clean(b.get("review_evidence_sha256")) == review_sha for b in batches if isinstance(b, dict)):
        summary = next((b for b in batches if isinstance(b, dict) and b.get("sha256") == sha and clean(b.get("review_evidence_sha256")) == review_sha), {})
        return out, {**summary, "idempotent_reuse": True}

    # Freeze the pre-ingest state for discovery diagnostics.  ``out`` is mutated as
    # records are admitted, so comparing later records against ``out`` would falsely
    # report same-batch manual duplicates as if automation had already discovered them.
    comparison_state = deepcopy(out)
    diagnostics: list[dict[str, Any]] = []
    recovery_queue: list[dict[str, Any]] = []
    weak_signal_recovery_queue: list[dict[str, Any]] = []
    counts = {
        "candidates": len(records), "found_in_corpus": 0, "scanner_seen_url_not_admitted": 0,
        "not_found": 0, "duplicate_in_batch": 0, "manual_admitted": 0, "deferred": 0,
        "rejected_core_gate": 0, "forthcoming": 0, "secondary_reference": 0,
        "context_only": 0, "weak_signal_candidates": 0, "manual_signals_admitted": 0,
        "verification_required": 0,
    }
    reviewed_resolved_ids: set[str] = set()
    admitted_item_ids: set[str] = set()
    matrix_item_ids: set[str] = set()
    for original_rec in records:
        review = review_evidence.get(clean(original_rec.get("manual_id")))
        rec = _record_with_review(original_rec, review)
        if rec.get("manual_candidate_kind") == "weak_signal":
            counts["weak_signal_candidates"] += 1
        if rec.get("manual_verification_required"):
            counts["verification_required"] += 1
        prior_strand_key, prior_matched, _prior_score = match_existing(rec, comparison_state)
        strand_key, matched, score = match_existing(rec, out)
        seen = _seen_automatically(rec, comparison_state)
        if prior_matched:
            automated_status = "found_in_corpus"; counts["found_in_corpus"] += 1
        elif seen:
            automated_status = "scanner_seen_url_not_admitted"; counts["scanner_seen_url_not_admitted"] += 1
        else:
            automated_status = "not_found"; counts["not_found"] += 1
        if matched:
            _provenance_value(matched, rec.get("manual_id", ""))
            if not prior_matched:
                counts["duplicate_in_batch"] += 1

        status = rec.get("manual_record_status")
        if status == "forthcoming_unpublished":
            counts["forthcoming"] += 1
            # A review pack may verify the bibliographic/source page, but the record remains
            # forthcoming and cannot become published evidence merely because text exists.
            reviewed = _retrieval_from_review(rec, review) if review else None
            retrieval = reviewed or {"retrieval_status": "not_attempted_forthcoming", "evidence_status": "forthcoming_unpublished", "text_mode": "metadata_only", "abstract": "", "body": ""}
            retrieval["evidence_status"] = "forthcoming_unpublished"
            ev = sr.gate_scope(rec.get("title", ""), clean(retrieval.get("abstract")), clean(retrieval.get("body")), _source_profile(rec)[1], source_kind=_source_profile(rec)[3])
            diagnostics.append(_diagnostic_record(rec, out, automated_status, retrieval, ev, "defer_forthcoming", matched, links_validated=links_validated, review=review, original_record=original_rec))
            continue
        if status == "context_outside_primary_window":
            counts["context_only"] += 1
            # Preserve reviewed text/provenance for context items while keeping them outside
            # the primary radar corpus and matrix.
            retrieval = (_retrieval_from_review(rec, review) if review else None) or {"retrieval_status": "not_attempted_context", "evidence_status": "context_reference", "text_mode": "metadata_only", "abstract": "", "body": ""}
            ev = sr.gate_scope(rec.get("title", ""), clean(retrieval.get("abstract")), clean(retrieval.get("body")), _source_profile(rec)[1], source_kind=_source_profile(rec)[3])
            diagnostics.append(_diagnostic_record(rec, out, automated_status, retrieval, ev, "retain_context_only", matched, links_validated=links_validated, review=review, original_record=original_rec))
            continue

        reviewed_retrieval = _retrieval_from_review(rec, review) if not matched else None
        retrieval = reviewed_retrieval or (retrieve_source(rec, session=session) if fetch and not matched else {
            "retrieval_status": "not_needed_existing_corpus" if matched else "not_attempted_offline",
            "evidence_status": "verified_existing_corpus" if matched else ("secondary_reference" if rec.get("manual_secondary_hint") else "uncertain_record"),
            "text_mode": "existing_corpus" if matched else "metadata_only", "abstract": "", "body": "", "resolved_url": rec.get("url", ""),
        })
        if rec.get("manual_secondary_hint") and not (review and review.get("resolved_primary")):
            retrieval["evidence_status"] = "secondary_reference"
        if retrieval.get("evidence_status") == "secondary_reference":
            counts["secondary_reference"] += 1

        source_name, tier, tier_label, source_kind = _source_profile(rec)
        abstract = clean(retrieval.get("abstract"))
        body = clean(retrieval.get("body"))
        # Existing corpus already passed the scanner's then-current gate; do not pretend the
        # manual citation text is new evidence.  Reuse its admitted status only for comparison.
        if matched:
            ev = {"a_pass": strand_key in {"strand_a", "frontier_evidence"}, "b_pass": strand_key == "strand_b", "eu_relevance": matched.get("eu_relevance"), "aboutness_reason": "existing_corpus", "text_mode": "existing_corpus", "ri_evidence": [], "geo_evidence": []}
            decision = "matched_existing_no_duplicate" if prior_matched else "matched_manual_batch_no_duplicate"
        else:
            ev = sr.gate_scope(rec.get("title", ""), abstract, body, tier, source_kind=source_kind)
            if _review_core_gate_verified(rec, review, retrieval):
                ev = _apply_review_core_gate(ev, review, retrieval)
            elif _review_core_gate_failed(rec, review):
                ev = _apply_review_core_gate_fail(ev, review)
            publication_date = _effective_publication_date(rec, retrieval)
            verified = retrieval.get("evidence_status") == "verified_primary_source" and bool(publication_date)
            event_verified = retrieval.get("evidence_status") in {"verified_primary_source", "verified_corroborated_current_event"} and bool(publication_date)
            decision = "defer_insufficient_or_unverified"
            if rec.get("manual_candidate_kind") == "weak_signal":
                # Manual weak signals never bypass the normal factual-news + Strand-A
                # anchoring route, and the substantive EU/European R&I gate remains a
                # prerequisite as an additional precision safeguard.
                if event_verified and ev.get("a_pass"):
                    signal = _new_public_signal(rec, retrieval, out, ingested_at, publication_date=publication_date, review=review)
                    if signal and _append_signal_if_new(out, signal):
                        counts["manual_signals_admitted"] += 1
                        decision = "admitted_verified_manual_weak_signal"
                        reviewed_resolved_ids.add(clean(rec.get("manual_id")))
                    elif signal:
                        decision = "matched_existing_no_duplicate"
                if decision.startswith("defer"):
                    counts["deferred"] += 1
            elif verified and (ev.get("a_pass") or ev.get("b_pass")):
                strands = []
                if ev.get("a_pass"): strands.append(("strand_a", "A"))
                if ev.get("b_pass"): strands.append(("strand_b", "B"))
                for dest, strand in strands:
                    if _append_if_new(out, dest, _new_public_item(rec, retrieval, ev, strand, ingested_at, publication_date=publication_date, review=review)):
                        counts["manual_admitted"] += 1
                        admitted_item_ids.add(clean(rec.get("manual_id")))
                        if review and review.get("matrix_evidence_verified") and _matrix_fields_from_review(rec, review):
                            matrix_item_ids.add(clean(rec.get("manual_id")))
                        reviewed_resolved_ids.add(clean(rec.get("manual_id")))
                decision = "admitted_verified_manual_source"
            elif _review_core_gate_failed(rec, review) and verified:
                # A verified underlying source that fails the substantive EU-R&I/geopolitics
                # gate is a real rejection, not a retrieval defer and not a recovery item.
                counts["rejected_core_gate"] += 1
                decision = "rejected_core_gate"
            else:
                counts["deferred"] += 1

        if matched:
            reviewed_resolved_ids.add(clean(rec.get("manual_id")))
            if review and review.get("matrix_evidence_verified"):
                matrix_fields = _matrix_fields_from_review(rec, review)
                matched.update(matrix_fields)
                if matrix_fields:
                    matrix_item_ids.add(clean(rec.get("manual_id")))
                claim = clean(review.get("display_claim"))
                if claim:
                    matched["core_message"] = sr.plain_language_claim(
                        matched.get("summary", ""), matched.get("title") or matched.get("headline") or "", claim
                    )
        diag = _diagnostic_record(rec, out, automated_status, retrieval, ev, decision, matched, links_validated=links_validated, review=review, original_record=original_rec)
        diag["source_tier"] = tier_label
        diag["source_kind"] = source_kind
        diagnostics.append(diag)
        # Exact-URL recovery is deliberately narrow. It improves recall without lowering
        # the substantive gate or adding whole low-precision domains.
        if (
            not matched and decision.startswith("defer") and rec.get("url")
            and not _is_generic_homepage(rec.get("url")) and rec.get("manual_record_status") == "candidate"
            and retrieval.get("evidence_status") not in {"secondary_reference", "forthcoming_unpublished"}
        ):
            queued = {
                "manual_id": rec.get("manual_id"), "title": rec.get("title"), "url": rec.get("url"),
                "date": rec.get("date"), "source": source_name, "tier": tier, "source_kind": source_kind,
                "date_precision": rec.get("date_precision"),
                "manual_verification_required": bool(rec.get("manual_verification_required")),
                "manual_link_status": "user_validated_reachable" if links_validated else "supplied_unchecked",
                "manual_candidate_kind": rec.get("manual_candidate_kind", "substantive"),
                "reason": diag.get("miss_reason"), "batch_id": batch_id,
            }
            if rec.get("manual_candidate_kind") == "weak_signal":
                weak_signal_recovery_queue.append(queued)
            else:
                recovery_queue.append(queued)

    # Merge queue/history rather than replacing prior batches; completed corpus matches are dropped.
    old_records = manual.get("records") if isinstance(manual.get("records"), list) else []
    old_queue = manual.get("recovery_queue") if isinstance(manual.get("recovery_queue"), list) else []
    old_signal_queue = manual.get("weak_signal_recovery_queue") if isinstance(manual.get("weak_signal_recovery_queue"), list) else []
    q_by_url = {_canonical_url(x.get("url", "")): x for x in old_queue + recovery_queue if isinstance(x, dict) and x.get("url") and not _is_generic_homepage(x.get("url"))}
    sq_by_url = {_canonical_url(x.get("url", "")): x for x in old_signal_queue + weak_signal_recovery_queue if isinstance(x, dict) and x.get("url") and not _is_generic_homepage(x.get("url"))}
    # If a queued URL is now in the corpus it no longer needs scanner recovery.
    corpus_urls = {_canonical_url(i.get("link", "")) for _, i in _corpus_entries(out)}
    merged_queue = [x for u, x in q_by_url.items() if u and u not in corpus_urls and clean(x.get("manual_id")) not in reviewed_resolved_ids]
    merged_signal_queue = [x for u, x in sq_by_url.items() if u and u not in corpus_urls and clean(x.get("manual_id")) not in reviewed_resolved_ids]
    counts.update({
        "automated_discovery_misses": counts["not_found"],
        "automated_seen_not_admitted": counts["scanner_seen_url_not_admitted"],
        "duplicates_or_already_admitted": counts["found_in_corpus"] + counts["duplicate_in_batch"],
        "newly_admitted_substantive_items": len(admitted_item_ids),
        "newly_admitted_matrix_items": len(matrix_item_ids & admitted_item_ids),
        "reviewed_matrix_items": len(matrix_item_ids),
    })
    batch = {
        "batch_id": batch_id, "source_file": source.name, "sha256": sha, "ingested_at": ingested_at,
        "fetch_attempted": bool(fetch or review_evidence), "runtime_fetch_enabled": bool(fetch),
        "review_evidence_items": len(review_evidence), "links_user_validated": bool(links_validated), "counts": counts,
        "review_evidence_sha256": review_sha, "reviewed_items": len(review_evidence),
    }
    batches = [b for b in batches if not (isinstance(b, dict) and b.get("sha256") == sha and clean(b.get("review_evidence_sha256")) == review_sha)] + [batch]
    manual.update({
        "profile_version": PROFILE_VERSION,
        "last_ingested_at": ingested_at,
        "batches": batches[-30:],
        "records": _merge_record_history(old_records, [{**d, "source_file": source.name, "batch_id": batch_id} for d in diagnostics]),
        "recovery_queue": merged_queue,
        "weak_signal_recovery_queue": merged_signal_queue,
        "recall_diagnostics": {
            "latest_batch_id": batch_id,
            "high_quality_discovery_misses": sum(1 for d in diagnostics if d.get("miss_reason") == "discovery_recall_gap"),
            "seen_but_not_admitted": sum(1 for d in diagnostics if d.get("automated_status") == "scanner_seen_url_not_admitted"),
            "direct_source_sitemap_unavailable": sum(1 for d in diagnostics if d.get("miss_reason") == "direct_source_sitemap_unavailable_in_saved_scan"),
            "bounded_rotation_not_observed": sum(1 for d in diagnostics if d.get("miss_reason") == "not_observed_during_bounded_direct_source_rotation"),
            "source_not_observed": sum(1 for d in diagnostics if d.get("miss_reason") == "source_not_observed_in_saved_direct_discovery_ledger"),
            "retrieval_insufficient": sum(1 for d in diagnostics if d.get("miss_reason") == "underlying_source_retrieval_failed_or_insufficient"),
            "exact_url_recovery_queue": len(merged_queue),
            "weak_signal_recovery_queue": len(merged_signal_queue),
            "failure_categories": dict(sorted(__import__("collections").Counter(
                clean(d.get("recall_failure_category")) for d in diagnostics if clean(d.get("recall_failure_category"))
            ).items())),
        },
    })
    out["manual_ingest"] = manual
    out["manual_ingest_profile_version"] = PROFILE_VERSION
    # Manual ingestion is not a live scan.  Preserve all scan timestamps/scan_result bookkeeping.
    out["last_updated"] = old_last_updated
    sr.normalize_reader_claims(out)
    summary = {**batch, "recovery_queue": len(merged_queue), "weak_signal_recovery_queue": len(merged_signal_queue), "idempotent_reuse": False}
    return out, summary


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Ingest a curated candidate/recovery list into radar state")
    ap.add_argument("manual_file", help="DOCX, PDF, CSV, JSON, YAML, TXT or Markdown candidate list")
    ap.add_argument("--state", default=str(sr.OUT_PATH), help="Input radar JSON (default: repository radar.json)")
    ap.add_argument("--out", default="", help="Output radar JSON (default: overwrite --state atomically)")
    ap.add_argument("--no-fetch", action="store_true", help="Do not retrieve cited URLs; compare/register/recover only")
    ap.add_argument("--refresh", action="store_true", help="Reprocess a file hash already ingested")
    ap.add_argument("--links-validated", action="store_true", help="Record supplied URLs as user-tested/reachable without treating that as evidence verification")
    ap.add_argument("--review-evidence", default="", help="JSON review pack containing evidence from independently reviewed underlying sources; normal gates still apply")
    args = ap.parse_args(argv)
    path = Path(args.manual_file)
    state_path = Path(args.state)
    out_path = Path(args.out) if args.out else state_path
    records = parse_manual_file(path)
    review_evidence = load_review_evidence(args.review_evidence) if args.review_evidence else {}
    state = json.loads(state_path.read_text(encoding="utf-8"))
    updated, summary = apply_manual_ingest(
        state, records, source_path=path, fetch=not args.no_fetch, refresh=args.refresh,
        links_validated=args.links_validated, review_evidence=review_evidence,
    )
    tmp = out_path.with_suffix(out_path.suffix + ".tmp")
    tmp.write_text(json.dumps(updated, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    tmp.replace(out_path)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
