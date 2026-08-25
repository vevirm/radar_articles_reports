import datetime as dt
import json
import subprocess
import sys
from pathlib import Path

from docx import Document

from scripts import manual_ingest as mi
from scripts import scan_radar as sr

ROOT = Path(__file__).resolve().parents[1]


class FakeResponse:
    def __init__(self, url, html, status=200, content_type='text/html'):
        self.url = url
        self.status_code = status
        self.text = html
        self.content = html.encode('utf-8')
        self.headers = {'content-type': content_type}


class FakeSession:
    def __init__(self, response):
        self.response = response
        self.calls = []

    def get(self, url, **kwargs):
        self.calls.append((url, kwargs))
        return self.response


def base_state():
    return {
        'last_updated': '2026-08-25T07:34Z',
        'first_scan_complete': True,
        'scan_results': {'new_a': 2, 'new_b': 0, 'new_c': 1},
        'scan_state': {'institution_seen_fingerprints': {}},
        'strand_a': [], 'strand_b': [], 'strand_c': [], 'frontier_evidence': [],
    }


def write_json_list(tmp_path, rows):
    p = tmp_path / 'manual.json'
    p.write_text(json.dumps(rows), encoding='utf-8')
    return p


def test_docx_manual_parser_handles_current_forthcoming_context_and_last_item(tmp_path):
    p = tmp_path / 'manual.docx'
    doc = Document()
    doc.add_paragraph('A. Talent and research security')
    doc.add_paragraph('A1 Example Institute (29 June 2026). European research talent under geopolitical competition. Example Institute.')
    doc.add_paragraph('https://example.org/report')
    doc.add_paragraph('Type 3 – Science-for-policy & academic. Curated candidate only.')
    doc.add_paragraph('A2 Example University (July 2026). EU research security and international collaboration. Example University.')
    doc.add_paragraph('https://example.edu/paper')
    doc.add_paragraph('4. Expected but not yet published')
    doc.add_paragraph('European Research Area Act – Expected Q3 2026.')
    doc.add_paragraph('5. Essential 2026 context')
    doc.add_paragraph('European Parliament, Resolution on research security, 10 March 2026.')
    doc.add_paragraph('6. Method')
    doc.save(p)
    rows = mi.parse_manual_file(p)
    assert [r['manual_id'] for r in rows[:2]] == ['A1', 'A2']
    assert rows[1]['title'].startswith('EU research security')
    assert any(r['manual_record_status'] == 'forthcoming_unpublished' for r in rows)
    assert any(r['manual_record_status'] == 'context_outside_primary_window' for r in rows)


def test_existing_match_gets_both_provenance_without_duplicate_or_scan_timestamp_change(tmp_path):
    st = base_state()
    st['strand_a'] = [{
        'title': 'European research security under geopolitical competition',
        'authors': 'A', 'source': 'Journal', 'date': '2026-06-01',
        'link': 'https://doi.org/10.1234/example', 'strand': 'A',
        'eu_relevance': 'direct', 'summary': 'Existing admitted evidence.',
        'core_message': 'Existing admitted evidence', 'new_this_scan': False,
    }]
    p = write_json_list(tmp_path, [{
        'id': 'M1', 'title': 'European research security under geopolitical competition',
        'url': 'https://doi.org/10.1234/example', 'date': '2026-06-01'
    }])
    rows = mi.parse_manual_file(p)
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, now=dt.datetime(2026, 8, 25, 8, 0, tzinfo=dt.timezone.utc))
    assert out['last_updated'] == st['last_updated']
    assert out['scan_results'] == st['scan_results']
    assert len(out['strand_a']) == 1
    assert out['strand_a'][0]['discovery_provenance'] == 'both'
    assert out['strand_a'][0]['provenance'] == ['automated_discovery', 'manual_candidate_ingestion']
    assert summary['counts']['found_in_corpus'] == 1


def test_metadata_only_manual_candidate_is_deferred_and_queued(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{
        'id': 'M1', 'title': 'European semiconductor research and economic security under geopolitical competition',
        'url': 'https://example.org/report', 'date': '2026-06-01'
    }])
    rows = mi.parse_manual_file(p)
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False)
    assert not out['strand_a'] and not out['strand_b']
    assert summary['counts']['deferred'] == 1
    rec = out['manual_ingest']['records'][0]
    assert rec['text_mode'] == 'metadata_only'
    assert rec['decision'] == 'defer_insufficient_or_unverified'
    assert len(out['manual_ingest']['recovery_queue']) == 1


def test_verified_abstract_runs_same_gate_and_can_be_admitted(tmp_path):
    st = base_state()
    title = 'European semiconductor research, strategic dependencies and competitiveness'
    abstract = (
        'The European Union is strengthening semiconductor research and innovation capacity as geopolitical competition and economic security concerns intensify. '
        'The analysis finds that dependence on non-EU fabrication and design-tool suppliers constrains European technological autonomy and can weaken competitiveness. '
        'It compares EU research infrastructure, strategic technology dependencies and innovation capability against the United States and China, and identifies supply resilience as a core policy challenge.'
    )
    html = f'<html><head><title>{title}</title><meta name="description" content="{abstract}"></head><body></body></html>'
    url = 'https://example.org/eu-semiconductor-report'
    p = write_json_list(tmp_path, [{'id': 'M1', 'title': title, 'url': url, 'date': '2026-06-01'}])
    rows = mi.parse_manual_file(p)
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=True, session=FakeSession(FakeResponse(url, html)))
    assert summary['counts']['manual_admitted'] >= 1
    assert out['strand_a']
    item = out['strand_a'][0]
    assert item['discovery_provenance'] == 'manual'
    assert item['evidence_status'] == 'verified_primary_source'
    assert item['source_text_mode'] == 'abstract_only'
    assert not item['new_this_scan']


def test_secondary_and_forthcoming_never_auto_admit(tmp_path):
    st = base_state()
    rows_data = [
        {'id': 'M1', 'title': 'Future EU defence R&D instruments', 'url': 'https://example.org/', 'note': 'Exact title to be confirmed; as reported by secondary source', 'date': '2026-07-30'},
        {'id': 'M2', 'title': 'European Research Area Act', 'status': 'forthcoming', 'date': '2026-09-01'},
    ]
    p = write_json_list(tmp_path, rows_data)
    rows = mi.parse_manual_file(p)
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False)
    assert not out['strand_a'] and not out['strand_b']
    assert summary['counts']['forthcoming'] == 1
    assert summary['counts']['secondary_reference'] == 1
    decisions = {r['manual_id']: r['decision'] for r in out['manual_ingest']['records']}
    assert decisions['M2'] == 'defer_forthcoming'


def test_manual_recovery_queue_is_bounded_and_does_not_lower_gate(monkeypatch):
    prev = {'manual_ingest': {'recovery_queue': [
        {'manual_id': f'M{i}', 'title': f'European semiconductor research and economic security under geopolitical competition {i}', 'url': f'https://example.org/r{i}', 'tier': 1}
        for i in range(20)
    ]}}
    monkeypatch.setattr(sr, 'KNOWN_AB_LINKS', {'https://example.org/r0'})
    jobs = sr.manual_recovery_jobs(prev, limit=4)
    assert len(jobs) == 4
    assert all(x['url'] != 'https://example.org/r0' for x in jobs)
    # The queue is discovery only: metadata alone still fails pass 1.
    ev = sr.gate_scope(jobs[0]['title'], '', '', 1, source_kind='institutional')
    assert not ev['a_pass']
    assert ev['aboutness_reason'] == 'insufficient_text'


def test_claimed_and_implied_quadrants_remain_distinct_in_frontier():
    script = r'''
global.RadarInsights=require('./briefing/insights.js');
const F=require('./frontier/frontier.js');
const d={strand_a:[{
  title:'EU compute dependence creates bottlenecks and competitiveness losses',
  source:'Test source', date:'2026-06-01', link:'https://example.org/x', strand:'A', eu_relevance:'direct',
  summary:'European Union AI infrastructure depends on non-EU cloud and chip suppliers. The dependency creates bottlenecks, shortages and competitiveness losses for European research and innovation.',
  core_message:'EU compute dependence creates bottlenecks and competitiveness losses',
  matrix_dimension:'infrastructure', quadrant_claimed:'A', quadrant_implied:'D'
}],strand_b:[],strand_c:[],frontier_evidence:[]};
const v=F.buildFrontier(d,{now:'2026-08-25T10:00:00+03:00'});
const x=v.signals[0];
console.log(JSON.stringify({column:x&&x.column.id,claimed:x&&x.quadrantClaimed,implied:x&&x.quadrantImplied}));
'''
    out = subprocess.run(['node', '-e', script], cwd=ROOT, text=True, capture_output=True, check=True)
    got = json.loads(out.stdout)
    assert got == {'column': 'D', 'claimed': 'A', 'implied': 'D'}


def test_csv_yaml_and_text_manual_formats_are_normalized(tmp_path):
    csv_p = tmp_path / 'manual.csv'
    csv_p.write_text('id,title,url,date,authors\nC1,European research security under geopolitical competition,https://example.org/c1,2026-06-12,Example Author\n', encoding='utf-8')
    csv_rows = mi.parse_manual_file(csv_p)
    assert csv_rows[0]['manual_id'] == 'C1'
    assert csv_rows[0]['title'].startswith('European research security')
    assert csv_rows[0]['date'] == '2026-06-12'

    yaml_p = tmp_path / 'manual.yaml'
    yaml_p.write_text('items:\n  - id: Y1\n    title: European innovation capacity in strategic competition\n    url: https://example.org/y1\n    date: 2026-07-01\n', encoding='utf-8')
    yaml_rows = mi.parse_manual_file(yaml_p)
    assert yaml_rows[0]['manual_id'] == 'Y1'
    assert yaml_rows[0]['url'] == 'https://example.org/y1'

    txt_p = tmp_path / 'manual.txt'
    txt_p.write_text('1. Candidates\nEuropean Commission (9 July 2026). European research and innovation under geopolitical competition. European Commission.\n', encoding='utf-8')
    txt_rows = mi.parse_manual_file(txt_p)
    assert txt_rows and txt_rows[0]['title'].startswith('European research and innovation')
    assert '.pdf' in mi.SUPPORTED


def test_verified_source_without_publication_date_is_not_admitted_or_fabricated(tmp_path):
    st = base_state()
    title = 'European semiconductor research and geopolitical competition'
    abstract = (
        'The European Union is strengthening semiconductor research and innovation capacity under geopolitical competition. '
        'The analysis finds non-EU dependencies constrain technological autonomy and European competitiveness.'
    )
    url = 'https://example.org/no-date'
    p = write_json_list(tmp_path, [{'id': 'M1', 'title': title, 'url': url}])
    rows = mi.parse_manual_file(p)
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=True, session=FakeSession(FakeResponse(url, f'<html><head><title>{title}</title><meta name="description" content="{abstract}"></head></html>')))
    assert summary['counts']['manual_admitted'] == 0
    assert not out['strand_a'] and not out['strand_b']
    assert out['manual_ingest']['records'][0]['decision'] == 'defer_insufficient_or_unverified'


def test_generic_homepage_is_not_put_in_exact_url_recovery_queue(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{'id': 'M1', 'title': 'European R&I policy analysis', 'url': 'https://example.org/', 'date': '2026-07-01'}])
    rows = mi.parse_manual_file(p)
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False)
    assert summary['counts']['deferred'] == 1
    assert out['manual_ingest']['recovery_queue'] == []


def test_frontier_numbered_docx_parser_preserves_candidates_weak_signals_and_curator_cells(tmp_path):
    p = tmp_path / 'frontier_additions.docx'
    doc = Document()
    doc.add_paragraph('EU R&I in Geopolitical Context — Additions')
    doc.add_paragraph('1. Knowledge & people')
    doc.add_paragraph('K1 Example Institute (23 June 2026). Europe as a research power under geopolitical competition. Example Institute.')
    doc.add_paragraph('https://example.org/k1')
    doc.add_paragraph('Type 4. Cells: K-A, K-D. Curator mapping only.')
    doc.add_paragraph('2. Infrastructure & inputs')
    doc.add_paragraph('I1 Example Institute (19 May 2026). Europe needs compute capacity for strategic research. Example Institute.')
    doc.add_paragraph('https://example.org/i1')
    doc.add_paragraph('Cells: I-B (primary); I-C.')
    doc.add_paragraph('5. Weak signals')
    doc.add_paragraph('W1 Example News (24 July 2026). Export controls disrupt European research inputs. Example News.')
    doc.add_paragraph('https://example.org/w1')
    doc.add_paragraph('Cells: I-D (primary); R-B.')
    doc.add_paragraph('6. Judgment calls and open items')
    doc.add_paragraph('This paragraph must not become a record.')
    doc.save(p)

    rows = mi.parse_manual_file(p)
    assert [r['manual_id'] for r in rows] == ['K1', 'I1', 'W1']
    assert rows[0]['curator_cells'] == ['K-A', 'K-D']
    assert rows[1]['curator_primary_cell'] == 'I-B'
    assert rows[1]['curator_cell_mapping_status'] == 'manual_hint_not_source_evidence'
    assert rows[2]['manual_candidate_kind'] == 'weak_signal'


def test_user_validated_links_are_recorded_but_do_not_bypass_evidence_or_date_verification(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{
        'id': 'M1',
        'title': 'European research security and strategic competition',
        'url': 'https://example.org/verified-link',
        'date': '2026',
        'note': 'Publication date within the window not yet confirmed.'
    }])
    rows = mi.parse_manual_file(p)
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, links_validated=True)
    assert summary['links_user_validated'] is True
    assert summary['fetch_attempted'] is False
    assert not out['strand_a'] and not out['strand_b']
    rec = out['manual_ingest']['records'][0]
    assert rec['manual_link_status'] == 'user_validated_reachable'
    assert rec['manual_verification_required'] is True
    assert rec['decision'] == 'defer_insufficient_or_unverified'
    assert out['manual_ingest']['recovery_queue'][0]['manual_link_status'] == 'user_validated_reachable'


def test_review_pack_can_supply_verified_primary_evidence_without_runtime_fetch(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{
        'id': 'K1',
        'title': 'Europe as a science power under geopolitical competition',
        'url': 'https://example.org/k1',
        'date': '2026-06-23',
        'note': 'Cells: K-A (primary); K-D.'
    }])
    rows = mi.parse_manual_file(p)
    review = {
        'K1': {
            'review_source_url': 'https://example.org/k1',
            'source_verified': True,
            'core_gate_verified': True,
            'primary_source': True,
            'published': '2026-06-23',
            'text_mode': 'abstract_only',
            'evidence_text': (
                'European Union research policy is responding to geopolitical competition with the United States and China. '
                'The source examines whether Europe can attract scientists and strengthen research and innovation capacity while '
                'funding and commercialization constraints limit its competitiveness.'
            ),
            'display_claim': 'Europe can attract research talent, but funding and innovation capacity constrain the scale of the gain.',
            'matrix_evidence_verified': True,
            'matrix_dimension': 'knowledge',
            'quadrant_implied': 'A',
            'quadrant_claimed': 'A',
            'matrix_basis': 'The underlying source links geopolitical disruption to European research-talent attraction and capacity.',
        }
    }
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_admitted'] == 1
    assert len(out['strand_a']) == 1
    item = out['strand_a'][0]
    assert item['discovery_provenance'] == 'manual'
    assert item['matrix_dimension'] == 'knowledge'
    assert item['quadrant_implied'] == 'A'
    assert item['quadrant_claimed'] == 'A'
    assert item['matrix_classification_source'] == 'reviewed_underlying_source'
    assert item['core_message'].startswith('Europe can attract research talent')
    assert not out['manual_ingest']['recovery_queue']


def test_review_pack_does_not_bypass_strict_core_gate(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{
        'id': 'K2',
        'title': 'European funder changes application rules after researcher feedback',
        'url': 'https://example.org/k2',
        'date': '2026-06-01',
        'note': 'Cells: K-A.'
    }])
    rows = mi.parse_manual_file(p)
    review = {
        'K2': {
            'source_verified': True,
            'primary_source': True,
            'published': '2026-06-01',
            'text_mode': 'abstract_only',
            'evidence_text': (
                'The European Research Council changed its resubmission rules after researchers objected to the proposed restrictions. '
                'The change concerns application eligibility, peer review workload and researcher feedback within the grant programme. '
                'The article focuses on application eligibility, peer review workload, grant rules and researcher feedback within the programme.'
            ),
            'matrix_evidence_verified': False,
        }
    }
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_admitted'] == 0
    assert not out['strand_a'] and not out['strand_b']
    rec = next(r for r in out['manual_ingest']['records'] if r['manual_id'] == 'K2')
    assert rec['decision'] == 'defer_insufficient_or_unverified'
    assert rec['source_review_status'] == ''


def test_reviewed_primary_resolution_can_replace_secondary_url_and_date(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{
        'id': 'C6',
        'title': 'European Tech Champions Initiative 2.0 — launch',
        'url': 'https://secondary.example/c6',
        'date': '2026',
        'note': 'Secondary source; substitute the primary source once located.'
    }])
    rows = mi.parse_manual_file(p)
    assert rows[0]['manual_secondary_hint'] is True
    review = {
        'C6': {
            'review_source_url': 'https://secondary.example/c6',
            'source_verified': True,
            'primary_source': True,
            'core_gate_verified': True,
            'review_status': 'reviewed_pass_core_gate',
            'resolved_primary': True,
            'resolved_url': 'https://primary.example/c6',
            'title': 'Europe launches 80 billion euro investment alliance to scale up tech leaders',
            'published': '2026-07-10',
            'source': 'European Investment Bank',
            'text_mode': 'abstract_only',
            'evidence_text': (
                'The European Investment Bank and all European Union member states launched a larger investment platform for highly innovative technology scale-ups. '
                'It addresses Europe’s scale-up financing gap so ideas, technologies and innovative firms born in Europe can stay and thrive in Europe, while reinforcing the European innovation ecosystem, strategic autonomy, innovation capacity, productivity and global competitiveness.'
            ),
            'display_claim': 'A pan-European scale-up fund is intended to keep technology champions and their value creation in Europe.',
            'matrix_evidence_verified': True,
            'matrix_dimension': 'conversion',
            'quadrant_implied': 'A',
        }
    }
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_admitted'] == 1
    item = out['strand_a'][0]
    assert item['link'] == 'https://primary.example/c6'
    assert item['date'] == '2026-07-10'
    assert item['source'] == 'European Investment Bank'
    assert item['manual_supplied_url'] == 'https://secondary.example/c6'
    assert item['review_resolved_url'] == 'https://primary.example/c6'
    assert out['manual_ingest']['recovery_queue'] == []
    rec = next(r for r in out['manual_ingest']['records'] if r['manual_id'] == 'C6')
    assert rec['review_url_bound_to_supplied_link'] is True


def test_review_pack_new_hash_reprocesses_same_manual_file_without_refresh(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{'id': 'M1', 'title': 'European research security', 'url': 'https://example.org/m1', 'date': '2026-07-01'}])
    rows = mi.parse_manual_file(p)
    out1, _ = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False)
    review = {'M1': {'source_verified': True, 'primary_source': True, 'published': '2026-07-01', 'evidence_text': 'European Union research security policy responds to geopolitical competition and protects research and innovation collaboration from strategic technology risks and foreign interference.', 'text_mode': 'abstract_only'}}
    out2, summary = mi.apply_manual_ingest(out1, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['idempotent_reuse'] is False
    assert summary['reviewed_items'] == 1


def test_reviewed_outside_window_correction_remains_context(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{'id': 'I4', 'title': 'Allied semiconductor export-control authorities', 'url': 'https://example.org/i4', 'date': '2026-05-26'}])
    rows = mi.parse_manual_file(p)
    review = {'I4': {'review_source_url': 'https://example.org/i4', 'source_verified': True, 'primary_source': True, 'published': '2025-03-14', 'record_status': 'context_outside_primary_window', 'evidence_text': 'The analysis compares allied legal authority for semiconductor export controls and geopolitical technology restrictions.', 'text_mode': 'abstract_only'}}
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['context_only'] == 1
    assert not out['strand_a'] and not out['strand_b']
    rec = next(r for r in out['manual_ingest']['records'] if r['manual_id'] == 'I4')
    assert rec['decision'] == 'retain_context_only'
    assert rec['reviewed_bibliographic_corrections']['date']['reviewed'] == '2025-03-14'


def test_reviewed_weak_signal_can_carry_explicit_matrix_evidence(tmp_path):
    st = base_state()
    # Anchor source must already exist because weak signals do not stand alone.
    st['strand_a'] = [{
        'title': 'EU-China research de-risking and dual-use export controls', 'link': 'https://example.org/a',
        'date': '2026-06-01', 'strand': 'A', 'eu_relevance': 'direct',
        'summary': 'European Union research and innovation policy addresses China-related technology dependencies, dual-use export controls and research de-risking under geopolitical competition.',
        'core_message': 'European R&I remains exposed to strategic input dependencies.'
    }]
    p = write_json_list(tmp_path, [{
        'id': 'W1', 'title': 'China restricts dual-use exports to EU entities', 'url': 'https://example.org/w1',
        'date': '2026-07-24', 'manual_candidate_kind': 'weak_signal', 'note': 'Cells: I-D (primary).'
    }])
    rows = mi.parse_manual_file(p)
    rows[0]['manual_candidate_kind'] = 'weak_signal'
    review = {'W1': {
        'review_source_url': 'https://example.org/w1',
        'source_verified': True, 'corroborated_current_event': True, 'published': '2026-07-24',
        'text_mode': 'abstract_only',
        'evidence_text': 'China imposed immediate dual-use export restrictions on named European Union entities in retaliation for EU sanctions. The affected entities include technology and research organisations, making the geopolitical measure a direct constraint on European research and innovation access to strategic dual-use inputs and cross-border technology supply.',
        'display_claim': 'China imposed immediate dual-use export restrictions on named EU entities in retaliation for EU sanctions.',
        'matrix_evidence_verified': True, 'matrix_dimension': 'infrastructure', 'quadrant_implied': 'D'
    }}
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_signals_admitted'] == 1
    signal = out['strand_c'][-1]
    assert signal['matrix_dimension'] == 'infrastructure'
    assert signal['quadrant_implied'] == 'D'


def test_review_pack_must_be_bound_to_exact_supplied_url(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{
        'id': 'K1', 'title': 'European research capacity under geopolitical competition',
        'url': 'https://example.org/supplied', 'date': '2026-06-23'
    }])
    rows = mi.parse_manual_file(p)
    review = {'K1': {
        'review_source_url': 'https://example.org/different',
        'source_verified': True, 'primary_source': True, 'core_gate_verified': True,
        'review_status': 'reviewed_pass_core_gate', 'published': '2026-06-23',
        'text_mode': 'abstract_only',
        'evidence_text': 'The European Union research and innovation system is responding to geopolitical competition and strategic technology dependencies with new capacity measures.'
    }}
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_admitted'] == 0
    rec = [x for x in out['manual_ingest']['records'] if x['manual_id'] == 'K1'][0]
    assert rec['review_url_bound_to_supplied_link'] is False
    assert rec['decision'] == 'defer_insufficient_or_unverified'


def test_reviewed_substantive_gate_is_not_blocked_by_keyword_heuristic(tmp_path):
    st = base_state()
    p = write_json_list(tmp_path, [{
        'id': 'C1', 'title': 'Exploring the investor landscape for venture capital',
        'url': 'https://example.org/ecb-vc', 'date': '2026-05-07'
    }])
    rows = mi.parse_manual_file(p)
    # Deliberately phrase the reviewed evidence without the scanner's strongest geopolitical
    # keywords. The separate reviewed adjudication establishes the substantive mechanism.
    review = {'C1': {
        'review_source_url': 'https://example.org/ecb-vc',
        'source_verified': True, 'primary_source': True, 'core_gate_verified': True,
        'review_status': 'reviewed_pass_core_gate', 'published': '2026-05-07',
        'text_mode': 'abstract_only',
        'evidence_text': ('The European Central Bank examines how limited European institutional-investor participation '
                          'constrains venture-capital fund scale and the ability of innovative firms to finance later-stage '
                          'growth. It links stronger European financing capacity to retaining value and reducing reliance on '
                          'outside capital in the innovation ecosystem.'),
        'display_claim': 'Weak institutional-investor participation constrains Europe’s scale-up financing capacity.',
        'matrix_evidence_verified': True, 'matrix_dimension': 'conversion',
        'quadrant_implied': 'C', 'quadrant_claimed': 'A',
        'matrix_basis': 'Current financing depends on outside capital; the advocated remedy is stronger home financing.'
    }}
    raw_ev = sr.gate_scope(rows[0]['title'], review['C1']['evidence_text'], '', 2, source_kind='institutional')
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_admitted'] == 1
    item = out['strand_a'][0]
    assert item['matrix_dimension'] == 'conversion'
    assert item['quadrant_implied'] == 'C'
    assert item['quadrant_claimed'] == 'A'
    assert item['link'] == 'https://example.org/ecb-vc'
    assert item['relevance_note']
    # The reviewed route is allowed even when lexical heuristics are narrower.
    assert item['source_review_status'] == 'reviewed_pass_core_gate'


def test_review_resolution_uses_primary_public_link_and_preserves_supplied_provenance(tmp_path):
    st = base_state()
    supplied = 'https://example.org/secondary-reference'
    p = write_json_list(tmp_path, [{
        'id': 'R2', 'title': 'European research security guidance', 'url': supplied,
        'date': '2026-07-09', 'note': 'secondary source; substitute primary source'
    }])
    rows = mi.parse_manual_file(p)
    review = {'R2': {
        'review_source_url': supplied, 'source_verified': True, 'primary_source': True,
        'resolved_primary': True, 'resolved_url': 'https://official.example.eu/primary',
        'core_gate_verified': True, 'review_status': 'reviewed_pass_core_gate',
        'published': '2026-07-09', 'text_mode': 'abstract_only',
        'evidence_text': ('The European Union and member-state research-security system applies screening and risk '
                          'management to international scientific collaboration in response to foreign interference and '
                          'technology-transfer concerns.'),
        'matrix_evidence_verified': True, 'matrix_dimension': 'rules', 'quadrant_implied': 'B',
        'matrix_basis': 'Stricter screening increases control while imposing collaboration costs.'
    }}
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_admitted'] == 1
    item = out['strand_a'][0]
    assert item['link'] == 'https://official.example.eu/primary'
    assert item['manual_supplied_url'] == supplied
    assert item['review_resolved_url'] == 'https://official.example.eu/primary'


def test_title_only_dedup_does_not_merge_distinct_tech_sovereignty_papers(tmp_path):
    st = base_state()
    rows_data = [
        {'id': 'R3', 'title': 'Does Europe Really Have a Plan for Tech Sovereignty?', 'url': 'https://example.org/r3', 'date': '2026-06-29'},
        {'id': 'R6', 'title': 'European Tech Sovereignty', 'url': 'https://example.org/r6', 'date': '2026-05-06'},
    ]
    p = write_json_list(tmp_path, rows_data)
    rows = mi.parse_manual_file(p)
    common = {
        'source_verified': True, 'primary_source': True, 'core_gate_verified': True,
        'review_status': 'reviewed_pass_core_gate', 'text_mode': 'abstract_only',
        'evidence_text': ('European Union technology and innovation policy addresses geopolitical dependence on foreign '
                          'technology suppliers and proposes measures to increase strategic autonomy and domestic capacity.'),
        'matrix_evidence_verified': True, 'matrix_dimension': 'rules', 'quadrant_implied': 'B',
        'matrix_basis': 'The source concerns protected European technological capacity.'
    }
    review = {
        'R3': {**common, 'review_source_url': 'https://example.org/r3', 'published': '2026-06-29'},
        'R6': {**common, 'review_source_url': 'https://example.org/r6', 'published': '2026-05-06'},
    }
    out, summary = mi.apply_manual_ingest(st, rows, source_path=p, fetch=False, review_evidence=review)
    assert summary['counts']['manual_admitted'] == 2
    assert len(out['strand_a']) == 2
    assert {x['manual_ingest_ids'][0] for x in out['strand_a']} == {'R3', 'R6'}


def test_reviewed_matrix_dimension_controls_row_but_curator_hint_does_not():
    script = r'''
global.RadarInsights=require('./briefing/insights.js');
const F=require('./frontier/frontier.js');
const d={strand_a:[{
  title:'EU science diplomacy framework', source:'Council', date:'2026-05-29', link:'https://example.org/r1',
  strand:'A', eu_relevance:'direct',
  summary:'The EU uses science cooperation as a foreign-policy instrument while balancing openness and research security.',
  core_message:'The EU adopted a common science-diplomacy framework that uses scientific cooperation as a foreign-policy tool.',
  matrix_dimension:'rules', quadrant_implied:'A', matrix_quadrant:'A',
  matrix_classification_source:'reviewed_underlying_source',
  matrix_evidence_basis:'The source establishes a common EU framework for the rules of international scientific engagement.',
  curator_primary_cell:'K-C'
}],strand_b:[],strand_c:[],frontier_evidence:[]};
const v=F.buildFrontier(d,{now:'2026-08-25T12:27:00+03:00'});
const x=v.signals[0];
console.log(JSON.stringify({row:x&&x.row.id,column:x&&x.column.id}));
'''
    out = subprocess.run(['node', '-e', script], cwd=ROOT, text=True, capture_output=True, check=True)
    assert json.loads(out.stdout) == {'row': 'rules', 'column': 'A'}


def test_reviewed_weak_signal_matrix_dimension_and_quadrant_are_honoured():
    script = r'''
global.RadarInsights=require('./briefing/insights.js');
const F=require('./frontier/frontier.js');
const d={strand_a:[{
  title:'EU China research security and strategic inputs', source:'Anchor', date:'2026-06-01', link:'https://example.org/a',
  strand:'A',eu_relevance:'direct',summary:'European Union research and innovation is exposed to China-related strategic input dependencies and export controls.',
  core_message:'European R&I is exposed to strategic input dependencies.'
}],strand_b:[],strand_c:[{
  headline:'China restricts dual-use exports to EU entities',title:'China restricts dual-use exports to EU entities',
  source:'News',date:'2026-07-24',link:'https://example.org/w1',anchor:'EU China research security and strategic inputs',watch_theme:'EU-China R&I de-risking',
  what:'China imposed immediate dual-use export restrictions on named EU technology and research organisations.',
  why_it_matters:'This directly constrains European access to strategic inputs.',
  matrix_dimension:'infrastructure',quadrant_implied:'D',matrix_quadrant:'D',
  matrix_classification_source:'reviewed_underlying_source',
  matrix_evidence_basis:'An external actor directly cuts strategic inputs to European entities.'
}],frontier_evidence:[]};
const v=F.buildFrontier(d,{now:'2026-08-25T12:27:00+03:00'});
const x=v.signals.find(s=>(s.bibliographicTitle||'').includes('restricts'));
console.log(JSON.stringify({row:x&&x.row.id,column:x&&x.column.id}));
'''
    out = subprocess.run(['node', '-e', script], cwd=ROOT, text=True, capture_output=True, check=True)
    assert json.loads(out.stdout) == {'row': 'infrastructure', 'column': 'D'}
